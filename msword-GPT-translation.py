"""
Translate MS Word documents into other languages by ChatGPT or GPT-4 API.
written by lancer1911
May 1, 2023
"""

import openai
import docx
import requests
import json
from configparser import ConfigParser
import time
import re
import argparse
import tkinter as tk
from tkinter import filedialog
import os

config = ConfigParser()
config.read("settings.cfg")

api_key = config.get("OpenAI", "api_key")
model = config.get("OpenAI", "model")
temperature = 0.2 

LANGUAGES = {
    "en": "English",
    "fr": "French",
    "de": "German",
    "jp": "Japanese",
    "es": "Spanish",
    "zh": "Simplified Chinese"
    # Add more languages if needed
}


def translate_text(text, source_lang, target_lang, custom_prompt, model, temperature):
    openai.api_key = api_key
    URL = config.get("ENDPOINTS", "URL")

    if model == "gpt-3.5-turbo":
        user_content = f"Translate the following {LANGUAGES[source_lang]} text to {LANGUAGES[target_lang]}: '{text}'"
    elif model == "gpt-4":
        user_content = f"{text}"
    else:
        raise ValueError("Invalid model specified in settings.cfg")

    payload = {
        "model": model,
        "temperature": temperature,
        "messages": [
            {"role": "system", "content": f"You are a professional translator particularly proficient in {custom_prompt}. You will translate {LANGUAGES[source_lang]} text to {LANGUAGES[target_lang]}. Please do not translate people's names, trademarks, and country codes. If you encounter any uncertain terms, please enclose the original text in parentheses after the translation."},
            {"role": "user", "content": user_content},
        ],
    }

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {openai.api_key}",
    }

    response = requests.post(URL, headers=headers, json=payload)
    response = response.json()

    print("\n\033[32mAPI response:\033[0m", response)

    if "error" in response and "Rate limit reached" in response["error"]["message"]:
        wait_time_str = response["error"]["message"].split("Please try again in ")[1].split(".")[0]
        wait_time = int(wait_time_str.rstrip("s"))
        print(f"\n\033[33mRate limit reached. Waiting for {wait_time} seconds.\033[0m")
        time.sleep(wait_time)
        return translate_text(text, source_lang, target_lang, custom_prompt, model, temperature)

    return response["choices"][0]["message"]["content"]

def translate_docx(input_path, output_path, source_lang, target_lang, custom_prompt, model, temperature):
    doc = docx.Document(input_path)
    translated_doc = docx.Document()

    # Set the default font to '宋体'
    font = translated_doc.styles['Normal'].font
    font.name = 'STSong'

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text:
            if paragraph.style.name.startswith('Heading'):
                # Preserve heading levels
                level = int(re.search(r'\d+', paragraph.style.name).group())
                translated_text = translate_text(text, source_lang, target_lang, custom_prompt, model, temperature)
                new_heading = translated_doc.add_heading(level=level)
                new_heading.text = translated_text

                # Set the font for heading to '宋体'
                new_heading.style.font.name = 'STSong'
            else:
                translated_text = translate_text(text, source_lang, target_lang, custom_prompt, model, temperature)
                new_para = translated_doc.add_paragraph(translated_text, style=paragraph.style)

                # Set the font for paragraph to '宋体'
                new_para.style.font.name = 'STSong'
        else:
            translated_doc.add_paragraph()

    translated_doc.save(output_path)

def get_input_file():
    parser = argparse.ArgumentParser(description="Translate a Word document using GPT.")
    parser.add_argument("input_file", nargs="?", help="Input .docx file to be translated")
    args = parser.parse_args()

    if args.input_file:
        return args.input_file
    else:
        root = tk.Tk()
        root.withdraw()
        return filedialog.askopenfilename(filetypes=[("Word documents", "*.docx")])

def get_output_file(input_path):
    file_name, file_ext = os.path.splitext(input_path)
    return f"{file_name}_translated{file_ext}"


if __name__ == "__main__":
    input_path = get_input_file() 
    output_path = get_output_file(input_path) 
    source_lang = config.get("Prompts", "source_lang")
    target_lang = config.get("Prompts", "target_lang")
    custom_prompt = config.get("Prompts", "custom_prompt")

    translate_docx(input_path, output_path, source_lang, target_lang, custom_prompt, model, temperature)
